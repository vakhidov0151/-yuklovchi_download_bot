import asyncio
import logging
import os
import uuid
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import CommandStart, Command
from aiogram.client.default import DefaultBotProperties
from aiogram.types import FSInputFile, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.client.session.aiohttp import AiohttpSession
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from config import BOT_TOKEN, GEMINI_API_KEY, ADMIN_IDS, CARD_NUMBER, CARD_OWNER
from downloader import get_video_info, download_media
from database import db
from lang import _
import io
import PIL.Image
import docx
import openpyxl
from fpdf import FPDF
import ai_summary
import html

logging.basicConfig(level=logging.INFO)

session = AiohttpSession(timeout=600)
bot = Bot(token=BOT_TOKEN, session=session, default=DefaultBotProperties(parse_mode='HTML'))
dp = Dispatcher()

from middleware import SubscriptionMiddleware
dp.message.middleware(SubscriptionMiddleware())
dp.callback_query.middleware(SubscriptionMiddleware())

url_cache = {}
text_cache = {}

class AdminStates(StatesGroup):
    waiting_for_broadcast = State()

class UserStates(StatesGroup):
    waiting_for_receipt = State()

import re

def sanitize_html(text: str) -> str:
    # Replace <p> with double newlines
    text = re.sub(r'<\/?p>', '\n\n', text)
    # Replace <br> with newline
    text = re.sub(r'<br\s*\/?>', '\n', text)
    # Replace <li> with bullet points
    text = re.sub(r'<li>', '• ', text)
    # Replace headings with bold
    text = re.sub(r'<h[1-6]>', '<b>', text)
    text = re.sub(r'<\/h[1-6]>', '</b>\n\n', text)
    
    # Strip all other tags except allowed ones
    allowed_tags = ['b', 'i', 'u', 's', 'a', 'code', 'pre', 'strong', 'em', 'ins', 'del', 'strike', 'blockquote']
    
    def replacer(match):
        tag = match.group(2).lower()
        if tag in allowed_tags:
            return match.group(0)
        return ""
        
    text = re.sub(r'<(\/?)([a-zA-Z0-9]+)[^>]*>', replacer, text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

async def send_long_message(wait_msg: types.Message, chat_id: int, text: str):
    if len(text) <= 4000:
        await wait_msg.edit_text(text)
        return
        
    await wait_msg.delete()
    
    paragraphs = text.split('\n\n')
    current_chunk = ""
    
    for p in paragraphs:
        if len(current_chunk) + len(p) + 2 > 4000:
            if current_chunk:
                try:
                    await bot.send_message(chat_id, current_chunk)
                except Exception:
                    # Fallback if HTML tags are broken across chunks
                    clean_chunk = re.sub(r'<[^>]+>', '', current_chunk)
                    await bot.send_message(chat_id, clean_chunk)
            current_chunk = p + "\n\n"
        else:
            current_chunk += p + "\n\n"
            
    if current_chunk.strip():
        try:
            await bot.send_message(chat_id, current_chunk)
        except Exception:
            clean_chunk = re.sub(r'<[^>]+>', '', current_chunk)
            await bot.send_message(chat_id, clean_chunk)

def get_lang_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇺🇿 O'zbekcha", callback_data="lang_uz")],
        [InlineKeyboardButton(text="🇷🇺 Русский", callback_data="lang_ru")],
        [InlineKeyboardButton(text="🇬🇧 English", callback_data="lang_en")]
    ])

@dp.message(CommandStart())
async def cmd_start(message: types.Message):
    args = message.text.split()
    referrer_id = None
    if len(args) > 1 and args[1].isdigit():
        referrer_id = int(args[1])
        if referrer_id == message.from_user.id:
            referrer_id = None
            
    is_new = db.add_user(
        telegram_id=message.from_user.id,
        full_name=message.from_user.full_name,
        username=message.from_user.username,
        referrer_id=referrer_id
    )
    
    if is_new and referrer_id:
        count = db.get_referral_count(referrer_id)
        try:
            await bot.send_message(referrer_id, f"🎉 Bitta do'stingiz botga qo'shildi! (Jami: {count} ta)")
            if count % 3 == 0:
                db.grant_pro(referrer_id, days=7)
                await bot.send_message(referrer_id, "🎁 <b>Tabriklaymiz!</b> Siz 3 ta do'stingizni taklif qildingiz va <b>7 kunlik PRO</b> tarifiga ega bo'ldingiz! Endi sizda limit yo'q!")
        except Exception:
            pass

    if is_new:
        await message.answer(_('uz', 'choose_lang'), reply_markup=get_lang_keyboard())
    else:
        lang = db.get_language(message.from_user.id)
        safe_name = html.escape(message.from_user.first_name)
        await message.answer(_(lang, 'start', name=safe_name))

@dp.message(Command("lang"))
async def cmd_lang(message: types.Message):
    await message.answer(_('uz', 'choose_lang'), reply_markup=get_lang_keyboard())

@dp.callback_query(F.data.startswith("lang_"))
async def process_lang(callback_query: types.CallbackQuery):
    lang_code = callback_query.data.split('_')[1]
    db.set_language(callback_query.from_user.id, lang_code)
    await callback_query.answer(_(lang_code, 'lang_changed'))
    safe_name = html.escape(callback_query.from_user.first_name)
    await callback_query.message.edit_text(_(lang_code, 'start', name=safe_name))

@dp.message(Command("stat"))
async def cmd_stat(message: types.Message):
    lang = db.get_language(message.from_user.id)
    users_count = db.count_users()
    await message.answer(_(lang, 'stat', count=users_count))

@dp.message(Command("profile"))
async def cmd_profile(message: types.Message):
    user_id = message.from_user.id
    lang = db.get_language(user_id)
    ref_count = db.get_referral_count(user_id)
    is_pro = db.is_pro(user_id)
    
    bot_info = await bot.get_me()
    ref_link = f"https://t.me/{bot_info.username}?start={user_id}"
    
    user_data = db.get_user(user_id)
    downloads_today = user_data[6] if user_data else 0
    
    if lang == 'uz':
        status = "💎 PRO (Cheksiz)" if is_pro else "🆓 FREE"
        limits = "Cheksiz" if is_pro else f"{5 - downloads_today} ta (kunlik)"
    elif lang == 'ru':
        status = "💎 PRO (Безлимит)" if is_pro else "🆓 FREE"
        limits = "Безлимитно" if is_pro else f"{5 - downloads_today} шт (дневной)"
    else:
        status = "💎 PRO (Unlimited)" if is_pro else "🆓 FREE"
        limits = "Unlimited" if is_pro else f"{5 - downloads_today} (daily)"
    
    await message.answer(_(lang, 'profile', status=status, limits=limits, refs=ref_count, link=ref_link))

@dp.message(Command("vip"))
async def cmd_vip(message: types.Message, state: FSMContext):
    lang = db.get_language(message.from_user.id)
    text = _(lang, 'vip_text', card=CARD_NUMBER, owner=CARD_OWNER)
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="cancel_vip")]
    ])
    await message.answer(text, reply_markup=keyboard)
    await state.set_state(UserStates.waiting_for_receipt)

@dp.callback_query(F.data == "cancel_vip")
async def cancel_vip(callback: types.CallbackQuery, state: FSMContext):
    await state.clear()
    await callback.message.delete()
    await callback.answer("Bekor qilindi")

@dp.message(UserStates.waiting_for_receipt, F.photo)
async def process_receipt(message: types.Message, state: FSMContext):
    await state.clear()
    if not ADMIN_IDS:
        await message.answer("Adminlar belgilanmagan. Iltimos keyinroq urinib ko'ring.")
        return
        
    user_id = message.from_user.id
    
    # Adminga jo'natish
    admin_text = f"💳 <b>Yangi to'lov!</b>\n\nFoydalanuvchi: <a href='tg://user?id={user_id}'>{message.from_user.first_name}</a> (ID: <code>{user_id}</code>)\nTarif: 1 Oylik PRO (15,000 so'm)\n\nTasdiqlaysizmi?"
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Tasdiqlash (1 Oy)", callback_data=f"approve_{user_id}")],
        [InlineKeyboardButton(text="❌ Rad etish", callback_data=f"reject_{user_id}")]
    ])
    
    await message.forward(chat_id=ADMIN_IDS[0])
    await bot.send_message(chat_id=ADMIN_IDS[0], text=admin_text, reply_markup=keyboard)
    
    await message.answer("⏳ Chekingiz adminga yuborildi. Tez orada tekshirib tasdiqlanadi va PRO tarifingiz yoqiladi!")

@dp.message(UserStates.waiting_for_receipt)
async def process_receipt_invalid(message: types.Message):
    await message.answer("Iltimos, to'lov skrinshotini (rasm) yuboring, yoki bekor qilish tugmasini bosing.")

@dp.callback_query(F.data.startswith("approve_"))
async def approve_vip(callback: types.CallbackQuery):
    if callback.from_user.id not in ADMIN_IDS:
        return
    
    user_id = int(callback.data.split("_")[1])
    db.grant_pro(user_id, days=30)
    
    lang = db.get_language(user_id)
    await bot.send_message(chat_id=user_id, text="🎉 Tabriklaymiz! To'lovingiz tasdiqlandi va sizga 30 kunlik PRO tarif yoqildi!")
    
    await callback.message.edit_text(callback.message.text + "\n\n<b>✅ TASDIQLANDI</b>", reply_markup=None)
    await callback.answer("Tasdiqlandi!")

@dp.callback_query(F.data.startswith("reject_"))
async def reject_vip(callback: types.CallbackQuery):
    if callback.from_user.id not in ADMIN_IDS:
        return
    
    user_id = int(callback.data.split("_")[1])
    
    await bot.send_message(chat_id=user_id, text="❌ Kechirasiz, siz yuborgan to'lov cheki tasdiqlanmadi. Agar xatolik bo'lsa adminga murojaat qiling.")
    
    await callback.message.edit_text(callback.message.text + "\n\n<b>❌ RAD ETILDI</b>", reply_markup=None)
    await callback.answer("Rad etildi!")

@dp.message(F.photo)
async def handle_photo(message: types.Message):
    lang = db.get_language(message.from_user.id)
    wait_msg = await message.answer(_(lang, 'wait_ai'))
    
    try:
        photo = message.photo[-1]
        file_info = await bot.get_file(photo.file_id)
        downloaded_file = await bot.download_file(file_info.file_path)
        img = PIL.Image.open(io.BytesIO(downloaded_file.read()))
        
        prompt = "Ushbu rasmdagi barcha yozuvlarni aniq va to'liq o'qib bering. Faqat rasmdagi tekstni o'zini bering. Agar rasmda jadval (table) bo'lsa, uni Excelga tushadigan qilib, qatorlarni yangi qator bilan, ustunlarni esa TAB (\\t) yoki '|' belgisi bilan ajratib yozing. Agar rasmda yozuv umuman bo'lmasa, nima tasvirlanganini qisqacha ta'riflab bering."
        if lang == 'ru':
            prompt = "Прочитайте весь текст на этом изображении точно и полностью. Верните только сам текст. Если на картинке есть таблица, отформатируйте ее для Excel: разделите столбцы знаком TAB или '|'. Если текста нет, кратко опишите, что изображено."
        elif lang == 'en':
            prompt = "Read all the text on this image accurately and completely. Return only the text itself. If there is a table in the image, format it for Excel by separating columns with a TAB or '|' character. If there is no text, briefly describe what is pictured."
            
        from google import genai
        from google.genai.errors import ServerError
        client = genai.Client(api_key=GEMINI_API_KEY)
        
        try:
            response = await client.aio.models.generate_content(
                model='gemini-3.6-flash',
                contents=[img, prompt]
            )
        except ServerError as e:
            if e.code == 503:
                response = await client.aio.models.generate_content(
                    model='gemini-3.5-flash',
                    contents=[img, prompt]
                )
            else:
                raise
        
        text_id = str(uuid.uuid4())[:8]
        text_cache[text_id] = response.text
        url_cache[text_id] = {'file_id': photo.file_id, 'type': 'photo'}
        
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(text="📄 Word (.docx)", callback_data=f"doc_docx_{text_id}"),
                InlineKeyboardButton(text="📝 Matn (.txt)", callback_data=f"doc_txt_{text_id}")
            ],
            [
                InlineKeyboardButton(text="📕 PDF (Matn)", callback_data=f"doc_pdf_{text_id}"),
                InlineKeyboardButton(text="📊 Excel (.xlsx)", callback_data=f"doc_xlsx_{text_id}")
            ],
            [
                InlineKeyboardButton(text="🖼 Rasmni PDF qilish", callback_data=f"img_pdf_{text_id}")
            ]
        ])
        
        display_text = response.text[:3000] + ("..." if len(response.text) > 3000 else "")
        await wait_msg.edit_text(_(lang, 'ai_result', text=display_text), reply_markup=keyboard)
        
    except Exception as e:
        await wait_msg.edit_text(f"❌ Error: {str(e)}")

import converter

@dp.message(F.video | F.audio | F.voice)
async def handle_media_summary(message: types.Message):
    lang = db.get_language(message.from_user.id)
    
    file_id = None
    if message.video: file_id = message.video.file_id
    elif message.audio: file_id = message.audio.file_id
    elif message.voice: file_id = message.voice.file_id
    
    if not file_id: return
    
    short_id = str(uuid.uuid4())[:8]
    url_cache[short_id] = {'file_id': file_id, 'type': 'media'}
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=_(lang, 'ai_summary_btn_doc'), callback_data=f"ai_sum_{short_id}")]
    ])
    
    await message.answer("🤖 Ushbu faylni eshitib/ko'rib, undagi gaplarni tushunib, sizga <b>qisqacha konspekt (xulosa)</b> yozib berishimni xohlaysizmi?", reply_markup=keyboard)


@dp.message(F.document)
async def handle_document(message: types.Message):
    lang = db.get_language(message.from_user.id)
    doc = message.document
    
    try:
        file_name = doc.file_name
        if not file_name:
            await message.answer("❌ Fayl nomi aniqlanmadi.")
            return
            
        if file_name.lower().endswith('.pdf'):
            short_id = str(uuid.uuid4())[:8]
            c_data = {'file_id': doc.file_id, 'orig_name': file_name}
            url_cache[short_id] = c_data
            db.set_cache(short_id, c_data)
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="📄 Word (.docx) ga o'tkazish", callback_data=f"c_p2d_{short_id}")],
                [InlineKeyboardButton(text="📊 Excel (.xlsx) ga o'tkazish", callback_data=f"c_p2x_{short_id}")]
            ])
            await message.answer("📕 Fayl formati: <b>PDF</b>\n\nQaysi formatga o'tkazamiz?", reply_markup=keyboard)
        elif file_name.lower().endswith('.docx'):
            short_id = str(uuid.uuid4())[:8]
            c_data = {'file_id': doc.file_id, 'orig_name': file_name}
            url_cache[short_id] = c_data
            db.set_cache(short_id, c_data)
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="📕 PDF ga o'tkazish", callback_data=f"c_d2p_{short_id}")]
            ])
            await message.answer("📄 Fayl formati: <b>Word (.docx)</b>\n\nQaysi formatga o'tkazamiz?", reply_markup=keyboard)
        else:
            await message.answer("❌ Kechirasiz, hozircha faqat <b>.pdf</b> va <b>.docx</b> formatidagi hujjatlarni qabul qilaman.")
            
    except Exception as e:
        import traceback
        err_msg = traceback.format_exc()
        import html
        await message.answer(f"❌ Xatolik yuz berdi:\n<pre>{html.escape(err_msg[-1000:])}</pre>")
        logging.error(f"Doc error: {err_msg}")

@dp.callback_query(F.data.startswith("ai_sum_"))
async def process_ai_summary(callback_query: types.CallbackQuery):
    lang = db.get_language(callback_query.from_user.id)
    short_id = callback_query.data.split('_')[2]
    
    cached = db.get_cache(short_id) or url_cache.get(short_id)
    if not cached or 'file_id' not in cached:
        await callback_query.answer("❌ Kesh topilmadi, faylni boshqatdan tashlang.", show_alert=True)
        return
        
    file_id = cached['file_id']
    
    await callback_query.answer()
    wait_msg = await callback_query.message.edit_text("⏳ <i>Media yuklab olinmoqda va AI tomonida tahlil qilinmoqda (bu 1-2 daqiqa olishi mumkin)... Kuting...</i>")
    
    try:
        file_info = await bot.get_file(file_id)
        
        # Max file size checking. Download can take long for very large files.
        # file_size is in bytes. Let's limit to 20MB just in case (Telegram limit for bot.get_file is exactly 20MB anyway)
        downloaded_file = await bot.download_file(file_info.file_path)
        
        os.makedirs("downloads", exist_ok=True)
        # We need a proper extension for Gemini. Try to guess from file_path or just use .mp4/.mp3
        ext = os.path.splitext(file_info.file_path)[1]
        if not ext: ext = ".mp4" # fallback
        local_path = f"downloads/ai_media_{short_id}{ext}"
        
        with open(local_path, 'wb') as f:
            f.write(downloaded_file.read())
            
        lang_str = "UZBEK"
        if lang == "ru":
            lang_str = "RUSSIAN"
        elif lang == "en":
            lang_str = "ENGLISH"
            
        prompt = (
            "Listen to this audio/video carefully. "
            "Write a detailed and structured summary of what is being said. "
            "Include key points, main topics discussed, and any important conclusions. "
            "If it's a song, summarize its lyrics/theme. "
            f"IMPORTANT: Your entire response MUST be in {lang_str} language. "
            "IMPORTANT: Format your response beautifully using ONLY HTML tags (<b> for bold, <i> for italic, <u> for underline). "
            "DO NOT use markdown asterisks (**) for bolding, use <b> tags instead! "
            "Do not wrap the whole response in ```html blocks, just return the raw HTML string."
        )
        
        summary = await asyncio.wait_for(ai_summary.summarize_media(local_path, prompt), timeout=600.0)
        
        full_text = f"🤖 <b>AI Xulosa:</b>\n\n{sanitize_html(summary)}"
        await send_long_message(wait_msg, wait_msg.chat.id, full_text)
        
        if os.path.exists(local_path): os.remove(local_path)
        
    except Exception as e:
        import traceback
        err_msg = traceback.format_exc()
        await wait_msg.edit_text(f"❌ Xatolik yuz berdi:\n<pre>{html.escape(err_msg[-1000:])}</pre>")
        logging.error(f"AI Summary Error: {err_msg}")


@dp.callback_query(F.data.startswith("c_"))
async def process_conversion(callback_query: types.CallbackQuery):
    lang = db.get_language(callback_query.from_user.id)
    data = callback_query.data.split('_')
    action = data[1] # p2d, p2x, d2p
    short_id = data[2]
    
    cached = db.get_cache(short_id) or url_cache.get(short_id)
    if not cached or 'file_id' not in cached:
        await callback_query.answer("❌ Kesh topilmadi, faylni boshqatdan tashlang.", show_alert=True)
        return
        
    file_id = cached['file_id']
    orig_name = cached.get('orig_name', 'Hujjat')
    base_name = orig_name.rsplit('.', 1)[0]
    
    await callback_query.answer()
    wait_msg = await callback_query.message.edit_text("⏳ <i>Fayl yuklab olinmoqda va konvertatsiya qilinmoqda... Kuting...</i>")
    
    try:
        file_info = await bot.get_file(file_id)
        downloaded_file = await bot.download_file(file_info.file_path)
        
        os.makedirs("downloads", exist_ok=True)
        in_path = f"downloads/in_{short_id}"
        out_path = f"downloads/out_{short_id}"
        
        if action == "p2d":
            in_path += ".pdf"
            out_path += ".docx"
            final_name = f"{base_name}.docx"
        elif action == "d2p":
            in_path += ".docx"
            out_path += ".pdf"
            final_name = f"{base_name}.pdf"
        elif action == "p2x":
            in_path += ".pdf"
            out_path += ".xlsx"
            final_name = f"{base_name}.xlsx"
            
        with open(in_path, 'wb') as f:
            f.write(downloaded_file.read())
            
        if action == "p2d":
            await converter.convert_pdf_to_docx(in_path, out_path)
        elif action == "d2p":
            await converter.convert_docx_to_pdf(in_path, out_path)
        elif action == "p2x":
            success = await converter.convert_pdf_to_xlsx(in_path, out_path)
            if not success:
                await wait_msg.edit_text("❌ Ushbu PDF ichida hech qanday jadval topilmadi! Excelga o'tkazish uchun jadval bo'lishi shart.")
                os.remove(in_path)
                return
                
        doc_file = FSInputFile(out_path, filename=final_name)
        await callback_query.message.answer_document(document=doc_file, caption=_(lang, 'doc_ready'))
        await wait_msg.delete()
        
        if os.path.exists(in_path): os.remove(in_path)
        if os.path.exists(out_path): os.remove(out_path)
        
    except Exception as e:
        import traceback
        err_msg = traceback.format_exc()
        await wait_msg.edit_text(f"❌ Xatolik yuz berdi:\n<pre>{html.escape(err_msg[-1000:])}</pre>")
        logging.error(f"Conv Error: {err_msg}")

@dp.callback_query(F.data.startswith("img_pdf_"))
async def process_img2pdf(callback_query: types.CallbackQuery):
    lang = db.get_language(callback_query.from_user.id)
    text_id = callback_query.data.split('_')[2]
    
    cache_data = url_cache.get(text_id)
    if not cache_data or 'file_id' not in cache_data:
        await callback_query.answer("❌ Error.", show_alert=True)
        return
        
    file_id = cache_data['file_id']
    
    await callback_query.answer()
    wait_msg = await callback_query.message.edit_text("⏳ <i>Rasm PDF ga aylantirilmoqda... Kuting...</i>")
    
    try:
        file_info = await bot.get_file(file_id)
        downloaded_file = await bot.download_file(file_info.file_path)
        
        os.makedirs("downloads", exist_ok=True)
        in_path = f"downloads/in_{file_id}.jpg"
        out_path = f"downloads/out_{file_id}.pdf"
        
        with open(in_path, 'wb') as f:
            f.write(downloaded_file.read())
            
        await converter.convert_image_to_pdf(in_path, out_path)
        
        doc_file = FSInputFile(out_path)
        await callback_query.message.answer_document(document=doc_file, caption=_(lang, 'doc_ready'))
        await wait_msg.delete()
        
        if os.path.exists(in_path): os.remove(in_path)
        if os.path.exists(out_path): os.remove(out_path)
        
    except Exception as e:
        await wait_msg.edit_text(f"❌ Xatolik yuz berdi: {e}")

@dp.callback_query(F.data.startswith("doc_"))
async def process_document(callback_query: types.CallbackQuery):
    lang = db.get_language(callback_query.from_user.id)
    data = callback_query.data.split('_')
    ext = data[1]
    text_id = data[2]
    
    content = text_cache.get(text_id)
    if not content:
        await callback_query.answer("❌ Error.", show_alert=True)
        return
        
    await callback_query.answer("Wait...")
    os.makedirs("downloads", exist_ok=True)
    file_name = f"downloads/AI_Natija_{text_id}.{ext}"
    
    try:
        if ext == "txt":
            with open(file_name, "w", encoding="utf-8") as f:
                f.write(content)
        elif ext == "docx":
            doc = docx.Document()
            doc.add_heading("AI Text", 0)
            doc.add_paragraph(content)
            doc.save(file_name)
        elif ext == "xlsx":
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Natija"
            lines = content.split('\n')
            current_row = 1
            for line in lines:
                if not line.strip() or line.strip().replace('-', '').replace('|', '') == '':
                    continue # Skip empty lines or markdown separator lines like |---|
                
                if '|' in line:
                    cols = [col.strip() for col in line.split('|')]
                    # Remove empty strings at start and end caused by markdown tables
                    if cols and not cols[0]: cols.pop(0)
                    if cols and not cols[-1]: cols.pop()
                elif '\t' in line:
                    cols = [col.strip() for col in line.split('\t')]
                else:
                    cols = [line.strip()]
                    
                for col_idx, val in enumerate(cols, start=1):
                    ws.cell(row=current_row, column=col_idx, value=val)
                current_row += 1
            wb.save(file_name)
        elif ext == "pdf":
            pdf = FPDF()
            pdf.add_page()
            if os.path.exists("Roboto-Regular.ttf"):
                pdf.add_font("Roboto", "", "Roboto-Regular.ttf")
                pdf.set_font("Roboto", size=12)
            else:
                pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, text=content)
            pdf.output(file_name)
            
        doc_file = FSInputFile(file_name)
        await callback_query.message.answer_document(
            document=doc_file, 
            caption=_(lang, 'doc_ready')
        )
    except Exception as e:
        await callback_query.message.answer(f"❌ Error: {e}")
    finally:
        if os.path.exists(file_name):
            os.remove(file_name)

@dp.callback_query(F.data.startswith("ai_sumurl_"))
async def process_ai_summary_url(callback_query: types.CallbackQuery):
    lang = db.get_language(callback_query.from_user.id)
    video_id = callback_query.data.split('_')[2]
    
    cached = db.get_cache(video_id) or url_cache.get(video_id)
    if not cached:
        await callback_query.answer("❌ Kesh topilmadi, havolani qaytadan yuboring.", show_alert=True)
        return
        
    url = cached['url']
    await callback_query.answer()
    
    wait_msg = await callback_query.message.answer("⏳ <i>YouTube/Instagram audiosi yuklab olinmoqda va AI tomonida tahlil qilinmoqda (bu bir necha daqiqa olishi mumkin)... Kuting...</i>")
    
    file_path = None
    try:
        # Faqat audioni yuklab olamiz (tezroq bo'lishi va hajm kichik bo'lishi uchun)
        file_path, _ = await asyncio.wait_for(download_media(url, media_type='audio'), timeout=180.0)
        
        lang_str = "UZBEK"
        if lang == "ru":
            lang_str = "RUSSIAN"
        elif lang == "en":
            lang_str = "ENGLISH"
            
        prompt = (
            "Listen to this audio carefully. "
            "Write a detailed and structured summary of what is being said. "
            "Include key points, main topics discussed, and any important conclusions. "
            "If it's a song, summarize its lyrics/theme. "
            f"IMPORTANT: Your entire response MUST be in {lang_str} language. "
            "IMPORTANT: Format your response beautifully using ONLY HTML tags (<b> for bold, <i> for italic, <u> for underline, <code> for code). "
            "DO NOT use markdown asterisks (**) for bolding, use <b> tags instead! "
            "Do not wrap the whole response in ```html blocks, just return the raw HTML string."
        )
        
        summary = await asyncio.wait_for(ai_summary.summarize_media(file_path, prompt), timeout=600.0)
        
        full_text = f"🤖 <b>AI Xulosa:</b>\n\n{sanitize_html(summary)}"
        await send_long_message(wait_msg, wait_msg.chat.id, full_text)
        
    except asyncio.TimeoutError:
        await wait_msg.edit_text("❌ Xatolik: Yuklab olish yoki analiz vaqti tugadi. Uzun videolar ko'p vaqt oladi.")
    except Exception as e:
        import traceback
        err_msg = traceback.format_exc()
        await wait_msg.edit_text(f"❌ Xatolik yuz berdi:\n<pre>{html.escape(err_msg[-1000:])}</pre>")
        logging.error(f"AI URL Summary Error: {err_msg}")
    finally:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)

@dp.callback_query(F.data.startswith("dl_"))
async def process_download(callback_query: types.CallbackQuery):
    user_id = callback_query.from_user.id
    lang = db.get_language(user_id)
    data = callback_query.data.split('_')
    action = data[1]
    video_id = data[2]
    cache_data = db.get_cache(video_id) or url_cache.get(video_id)
    
    if not cache_data:
        await callback_query.answer("❌ Error: Kesh topilmadi, havolani qaytadan yuboring.", show_alert=True)
        return
        
    url = cache_data['url']
    
    if action == "pic":
        thumb_url = cache_data.get('thumbnail')
        if thumb_url:
            await callback_query.message.answer_photo(photo=thumb_url)
            await callback_query.answer()
        return

    if not db.check_limit(user_id):
        await callback_query.answer("❌ Limit!", show_alert=True)
        await callback_query.message.answer(_(lang, 'limit_over'))
        return

    await callback_query.answer()
    
    # Text message bo'lsa edit qilamiz, video/rasm bo'lsa yangi xabar yozamiz
    if callback_query.message.text:
        wait_msg = await callback_query.message.edit_text(_(lang, 'downloading'))
    else:
        wait_msg = await callback_query.message.answer(_(lang, 'downloading'))
        
    try:
        if action == "vid":
            quality = data[3]
            file_path, title = await asyncio.wait_for(download_media(url, media_type='video', quality=quality), timeout=600.0)
            
            if os.path.getsize(file_path) > 49.5 * 1024 * 1024:
                if os.path.exists(file_path):
                    os.remove(file_path)
                await wait_msg.edit_text("❌ Ushbu video hajmi 50 MB dan katta. Telegram botlari orqali faqat 50 MB gacha bo'lgan fayllarni yuborish mumkin.\n\nIltimos, videoning pastroq sifatini (masalan 360p yoki 480p) tanlang yoki faqat Audio sifatida yuklab oling!")
                return
                
            clean_title = re.sub(r'[\\/*?:"<>|\n\r]', '', str(title))[:50].strip() or "video"
            video = FSInputFile(file_path, filename=f"{clean_title}.mp4")
            msg = await callback_query.message.answer_video(video=video, caption=f"🎬 <b>{html.escape(str(title)[:100])}</b>\n💿 {quality}p")
            
            # AI summary tugmasini qo'shish
            short_id = str(uuid.uuid4())[:8]
            url_cache[short_id] = {'file_id': msg.video.file_id, 'type': 'media'}
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text=_(lang, 'ai_summary_btn_doc'), callback_data=f"ai_sum_{short_id}")]
            ])
            await msg.edit_reply_markup(reply_markup=keyboard)
            
        elif action == "aud":
            file_path, title = await asyncio.wait_for(download_media(url, media_type='audio'), timeout=600.0)
            
            if os.path.getsize(file_path) > 49.5 * 1024 * 1024:
                if os.path.exists(file_path):
                    os.remove(file_path)
                await wait_msg.edit_text("❌ Ushbu audio hajmi 50 MB dan katta bo'lgani uchun Telegram orqali yuborish imkonsiz.")
                return
                
            clean_title = re.sub(r'[\\/*?:"<>|\n\r]', '', str(title))[:50].strip() or "audio"
            ext = os.path.splitext(file_path)[1] or ".m4a"
            audio = FSInputFile(file_path, filename=f"{clean_title}{ext}")
            msg = await callback_query.message.answer_audio(audio=audio, caption=f"🎵 <b>{html.escape(str(title)[:100])}</b>")
            
            # AI summary tugmasini qo'shish
            short_id = str(uuid.uuid4())[:8]
            url_cache[short_id] = {'file_id': msg.audio.file_id, 'type': 'media'}
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text=_(lang, 'ai_summary_btn_doc'), callback_data=f"ai_sum_{short_id}")]
            ])
            await msg.edit_reply_markup(reply_markup=keyboard)
            
        await wait_msg.delete()
        
        # Agar callback qilingan xabar text bo'lsa uni o'chiramiz, video bo'lsa uni o'chirmaymiz (chunki userga qolsin)
        if callback_query.message.text:
            await callback_query.message.delete()
            
        if os.path.exists(file_path):
            os.remove(file_path)
        db.add_download(user_id)
        
    except asyncio.TimeoutError:
        await wait_msg.edit_text("❌ Xatolik: Yuklab olish vaqti tugadi (10 daqiqadan oshib ketdi). VPN tezligi pastligi sabab bo'lishi mumkin.")
    except Exception as e:
        await wait_msg.edit_text(f"❌ Xatolik: {e}")

@dp.message(Command("admin"))
async def admin_panel(message: types.Message, state: FSMContext):
    if message.from_user.id not in ADMIN_IDS:
        return
        
    total_users = db.count_users()
    
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📢 Barchaga xabar yuborish", callback_data="admin_broadcast")]
    ])
    
    await message.answer(f"👨‍💻 <b>Admin Panel</b>\n\n👥 Umumiy foydalanuvchilar: <b>{total_users} ta</b>", reply_markup=keyboard)

@dp.callback_query(F.data == "admin_broadcast")
async def ask_broadcast_msg(callback: types.CallbackQuery, state: FSMContext):
    if callback.from_user.id not in ADMIN_IDS:
        return
        
    await callback.message.answer("Xabarni yuboring (Matn, rasm yoki video bo'lishi mumkin).\n\n<i>Bekor qilish uchun /cancel tugmasini bosing.</i>")
    await state.set_state(AdminStates.waiting_for_broadcast)
    await callback.answer()

@dp.message(Command("cancel"))
async def cancel_broadcast(message: types.Message, state: FSMContext):
    if message.from_user.id not in ADMIN_IDS:
        return
    await state.clear()
    await message.answer("❌ Xabar yuborish bekor qilindi.")

@dp.message(AdminStates.waiting_for_broadcast)
async def send_broadcast(message: types.Message, state: FSMContext):
    if message.from_user.id not in ADMIN_IDS:
        return
        
    await state.clear()
    users = db.get_all_users()
    
    if not users:
        await message.answer("Bazada foydalanuvchilar yo'q.")
        return
        
    wait_msg = await message.answer(f"⏳ Xabar {len(users)} ta foydalanuvchiga yuborilmoqda. Iltimos kuting...")
    
    success = 0
    fail = 0
    
    for user_id in users:
        try:
            await message.copy_to(chat_id=user_id)
            success += 1
        except Exception:
            fail += 1
            
        await asyncio.sleep(0.05) # Telegram spam limitlaridan saqlanish uchun
        
    await wait_msg.edit_text(f"✅ <b>Xabar tarqatish yakunlandi!</b>\n\n✔️ Muvaffaqiyatli: {success} ta\n❌ Yetib bormadi (botni bloklaganlar): {fail} ta")


@dp.message()
async def handle_message(message: types.Message):
    try:
        lang = db.get_language(message.from_user.id)
        text = message.text
        
        if not text:
            return
            
        text = text.strip()
        if text.startswith("http://") or text.startswith("https://"):
            # Instagram, TikTok, YouTube Shorts, X (Twitter), Threads, Facebook, Pinterest bo'lsa - DARHOL YUKLAB BERAMIZ!
            direct_domains = [
                'instagram.com', 'tiktok.com', '/shorts/', 
                'twitter.com', 'x.com', 'threads.net', 
                'facebook.com', 'fb.watch', 'pinterest.com', 'pin.it'
            ]
            is_direct = any(domain in text.lower() for domain in direct_domains)
            
            if is_direct:
                if not db.check_limit(message.from_user.id):
                    await message.answer(_(lang, 'limit_over'))
                    return
                    
                wait_msg = await message.answer("⏳ <i>Video yuklanmoqda... Kuting...</i>")
                try:
                    file_path, v_title = await asyncio.wait_for(download_media(text, media_type='video'), timeout=90.0)
                    if not file_path or not os.path.exists(file_path):
                        await wait_msg.edit_text("❌ Xatolik: Video yuklanmadi. Havolani tekshiring.")
                        return
                        
                    if os.path.getsize(file_path) > 49.5 * 1024 * 1024:
                        os.remove(file_path)
                        await wait_msg.edit_text("❌ Ushbu video hajmi 50 MB dan katta bo'lgani uchun Telegram orqali yuborib bo'lmadi.")
                        return
                        
                    clean_title = re.sub(r'[\\/*?:"<>|\n\r]', '', str(v_title))[:50].strip() or "video"
                    video = FSInputFile(file_path, filename=f"{clean_title}.mp4")
                    
                    video_id = str(uuid.uuid4())[:8]
                    c_item = {'url': text, 'thumbnail': None}
                    url_cache[video_id] = c_item
                    db.set_cache(video_id, c_item)
                    
                    keyboard = InlineKeyboardMarkup(inline_keyboard=[
                        [
                            InlineKeyboardButton(text="🎵 Musiqasi (Audio)", callback_data=f"dl_aud_{video_id}"),
                            InlineKeyboardButton(text="🧠 AI Konspekt", callback_data=f"ai_sumurl_{video_id}")
                        ]
                    ])
                    
                    await message.answer_video(video=video, caption=f"🎬 <b>{html.escape(str(v_title)[:100])}</b>", reply_markup=keyboard)
                    await wait_msg.delete()
                    
                    if os.path.exists(file_path):
                        os.remove(file_path)
                    db.add_download(message.from_user.id)
                    return
                except asyncio.TimeoutError:
                    await wait_msg.edit_text("❌ Tarmoq xatosi: Yuklab olish vaqti tugadi. Qaytadan urinib ko'ring.")
                    return
                except Exception as e:
                    await wait_msg.edit_text(f"❌ Xatolik: {e}")
                    return

            # 2. Agar uzun YouTube video bo'lsa - ma'lumotlarini olib, sifat tanlash menyusini chiqaramiz
            wait_msg = await message.answer(_(lang, 'wait_video'))
            
            try:
                info = await asyncio.wait_for(get_video_info(text), timeout=30.0)
            except asyncio.TimeoutError:
                await wait_msg.edit_text("❌ Tarmoq xatosi: Qidiruv vaqti tugadi. Server biroz band, iltimos qayta urinib ko'ring.")
                return
            except Exception as e:
                await wait_msg.edit_text(f"❌ Xatolik: {e}")
                return
                
            if not info:
                await wait_msg.edit_text(_(lang, 'not_found'))
                return

            title = info.get('title', 'Video')
            duration = info.get('duration', 0)
            extractor = info.get('extractor', 'Unknown')
            
            video_id = str(uuid.uuid4())[:8]
            c_item = {
                'url': text,
                'thumbnail': info.get('thumbnail')
            }
            url_cache[video_id] = c_item
            db.set_cache(video_id, c_item)
            
            keyboard = InlineKeyboardMarkup(inline_keyboard=[
                [
                    InlineKeyboardButton(text=_(lang, 'ai_summary_btn_audio'), callback_data=f"ai_sumurl_{video_id}")
                ],
                [
                    InlineKeyboardButton(text="⬇️ 1080p", callback_data=f"dl_vid_{video_id}_1080"),
                    InlineKeyboardButton(text="⬇️ 720p", callback_data=f"dl_vid_{video_id}_720")
                ],
                [
                    InlineKeyboardButton(text="⬇️ 480p", callback_data=f"dl_vid_{video_id}_480"),
                    InlineKeyboardButton(text="⬇️ 360p", callback_data=f"dl_vid_{video_id}_360")
                ],
                [
                    InlineKeyboardButton(text="🎵 Audio", callback_data=f"dl_aud_{video_id}"),
                    InlineKeyboardButton(text="🖼 Thumbnail", callback_data=f"dl_pic_{video_id}")
                ]
            ])
            
            mins, secs = divmod(duration or 0, 60)
            time_str = f"{mins}:{secs:02d}"
            
            await wait_msg.edit_text(_(lang, 'video_caption', title=title, ext=extractor, time=time_str), reply_markup=keyboard)
        else:
             await message.answer(_(lang, 'send_only_link'))
             
    except Exception as e:
        await message.answer(f"⚠️ Kritik xatolik yuz berdi:\n<code>{str(e)}</code>")

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n🛑 Bot to'xtatildi (Ctrl+C).")
